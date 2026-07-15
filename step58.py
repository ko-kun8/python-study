import os 
from docx import Document #新顔:word文書を操る大ボス

target_folder = "自動化テストフォルダ"
output_word_path = os.path.join(target_folder, "PYthonで作った書類.docx")

print("ーーー Wordファイルの作成処理を開始しますーーー")

try:
    # 1. 【新技】空っぽのWordドキュメントを新しく用意する
    doc = Document()

    # 2. 書類のタイトル（大きな文字）を追加する
    doc.add_heading(" 【自動化】業務報告書", level=0)

    # 3. 普通の文章（段落）を追加する
    p = doc.add_paragraph("お疲れ様です。Pythonのプログラムより自動生成された報告書です。")
    # 文末に太字で文字を付け足す
    p.add_run("(※この部分は太字の補足です)").bold = True

    # 4. 見出し（セクションの区切り）を追加する
    doc.add_heading("本日の実施内容", level=1)

    # 5. 箇条書き（リスト）を追加する
    doc.add_paragraph("osモジュールを使ったフォルダ操作の学習", style="List Bullet")
    doc.add_paragraph("Pillowを使った画像一括リサイズシステムの開発", style="List Bullet")
    doc.add_paragraph("python -docxを使ったWord書類の自動生成テスト", style="List Bullet")

    # 6. 【超重要】組み立てたデータをWordファイルとして保存
    doc.save(output_word_path)

    print(f" Wordファイルの作成が完了しました！→{output_word_path}")

except Exception as e:
    print(f"エラーが発生しました: {e}")
    